import streamlit as st
import pandas as pd
import os, io
from docx import Document
from datetime import datetime, date

# --- 1. CONFIGURACIÓN ---
st.set_page_config(page_title="D.S.P. Honduras - Protocolo Completo", layout="wide")

# --- 2. MOTOR DE ANÁLISIS CLÍNICO CON JUSTIFICACIÓN TÉCNICA ---
def motor_ia_dsp(datos):
    cuerpo = f"{datos.get('motivo', '')} {datos.get('opinion', '')}".lower().strip()
    checks = datos.get('checks', [])
    
    if len(cuerpo) < 10 and not checks:
        return {
            "estado": "PENDIENTE",
            "rec": "Información insuficiente.",
            "just_rec": "Se requiere mayor exploración clínica para emitir recomendaciones responsables.",
            "test": "N/A",
            "just_test": "La selección de batería psicométrica depende de la hipótesis diagnóstica."
        }

    # Riesgo Autolítico
    if any(x in cuerpo for x in ["suicid", "muerte", "matar", "ganas de morir"]) or "Intentos suicidas" in checks or "Ganas de morir" in checks:
        return {
            "estado": "ALERTA: Riesgo Autolítico Detectado",
            "rec": "Remisión inmediata a Psiquiatría y activación de protocolo de vigilancia.",
            "just_rec": "La presencia de indicadores de muerte requiere intervención multidisciplinaria para preservar la integridad física del paciente y estabilización neuroquímica.",
            "test": "Inventario de Desesperanza de Beck (BHS) y Escala de Ideación Suicida de Beck.",
            "just_test": "Estos reactivos cuantifican objetivamente el nivel de desesperanza y la letalidad de la ideación, fundamentales para el triage clínico."
        }
    
    # Organicidad / Psicosis
    elif any(x in cuerpo for x in ["voces", "alucinacion", "extrañas", "convulsion"]) or "Escucha voces" in checks:
        return {
            "estado": "INDICADOR: Posible Compromiso Orgánico o Psicótico",
            "rec": "Interconsulta con Neurología y evaluación psiquiátrica.",
            "just_rec": "Es necesario descartar etiología orgánica (lesiones, tumores o disfunciones neurológicas) antes de concluir un diagnóstico meramente psicológico.",
            "test": "Test Gestáltico Visomotor de Bender y SCL-90-R.",
            "just_test": "El Bender evalúa la función gestáltica visomotora asociada a daño orgánico, y el SCL-90-R mapea dimensiones de psicoticismo y paranoia."
        }

    # Estrés Policial
    else:
        return {
            "estado": "ESTADO: Reacción de Ajuste / Estrés Laboral",
            "rec": "Entrenamiento en Higiene Mental y Terapia Breve de Apoyo.",
            "just_rec": "El personal policial está expuesto a estrés crónico; fortalecer los mecanismos de afrontamiento previene el escalonamiento a trastornos de ansiedad o burnout.",
            "test": "Cuestionario de Personalidad 16PF-5 y Test Proyectivo HTP.",
            "just_test": "El 16PF permite conocer la estructura de personalidad para manejar el estrés, y el HTP revela la percepción del yo y su entorno familiar."
        }

# --- 3. INTERFAZ: REPLICACIÓN EXACTA DEL PDF ---
st.title("📋 Entrevista Psicológica para Adultos - D.S.P.")
st.caption("Formulario Oficial Estandarizado - Dirección de Sanidad Policial")

tabs = st.tabs([
    "I-II. Identificación y Motivo", 
    "III-V. Salud y Síntomas", 
    "VI. Historia Familiar", 
    "VII-IX. Social y Hábitos", 
    "X. Desarrollo y Educación", 
    "XI-XII. Personalidad y Cierre"
])

# PÁGINA 1: DATOS Y MOTIVO
with tabs[0]:
    st.subheader("I. DATOS GENERALES")
    c1, c2, c3 = st.columns(3)
    nombre = c1.text_input("Nombre Completo:")
    lugar_f_nac = c2.text_input("Lugar y Fecha de Nacimiento:")
    nacionalidad = c3.text_input("Nacionalidad:", value="Hondureña")
    
    c4, c5, c6 = st.columns(3)
    sexo = c4.selectbox("Sexo:", ["M", "F"])
    est_civil = c5.text_input("Estado Civil:")
    edad = c6.text_input("Edad:")
    
    c7, c8, c9 = st.columns(3)
    religion = c7.text_input("Religión:")
    celular = c8.text_input("Celular:")
    ocupacion = c9.text_input("Ocupación actual:")
    
    c10, c11 = st.columns(2)
    asignacion = c10.text_input("Asignación:")
    militar = c11.radio("¿Prestó servicio militar?", ["No", "Sí"])
    
    c12, c13 = st.columns(2)
    direccion = c12.text_input("Dirección Actual:")
    nivel_edu = c13.text_input("Nivel educativo:")
    
    c14, c15, c16 = st.columns(3)
    remitido = c14.text_input("Remitido por:")
    pasatiempo = c15.text_input("Pasatiempo:")
    deportes = c16.text_input("Deportes:")

    st.subheader("II. MOTIVO DE CONSULTA")
    motivo = st.text_area("Escriba el motivo de la consulta:")

# PÁGINA 1-2: ANTECEDENTES Y SÍNTOMAS
with tabs[1]:
    st.subheader("III. ANTECEDENTES CLÍNICOS Y PSICOLÓGICOS")
    ant_sit = st.text_area("ANTECEDENTES DE LA SITUACIÓN (¿Cuándo se sintió bien la última vez? ¿Cómo se han desarrollado los síntomas? ¿Antecedentes familiares?)")
    funciones = st.text_input("Funciones orgánicas principales: sueño, apetito, sed y defecación.")
    
    c17, c18 = st.columns(2)
    alergias = c17.text_input("¿Padece usted alergias? (Especifique cuales)")
    meds = c18.text_input("¿Toma algún medicamento regularmente? (¿Para qué?)")
    
    enf_inf = st.text_input("¿Cuáles son las enfermedades que sufrió durante la infancia?")
    cirugias = st.text_input("¿Ha sido intervenido quirúrgicamente alguna vez? (Especifique)")
    hosp = st.text_input("¿Lo han hospitalizado? (¿Por qué?)")

    st.subheader("IV-V. SÍNTOMAS (Marque los presentes en su vida)")
    sintomas_list = [
        "Pesadillas", "Sonambulismo", "Terror nocturno", "Enuresis (orinarse)", "Encopresis (defecar)", 
        "Onicofagia (uñas)", "Tics nerviosos", "Fobias", "Drogas", "Alcohol", "Ideas suicidas", 
        "Intentos suicidas", "Alucinaciones", "Pérdida de memoria", "Mareos", "Desmayos", 
        "Accidentes", "Maltrato Físico", "Tartamudez", "Escucha voces", "Caminar dormido", 
        "Miedos", "Cólico/Diarrea tensional", "Golpes en la cabeza", "Hablar dormido", 
        "Ver cosas extrañas", "Convulsiones", "Ganas de morir", "Fiebre", 
        "Problemas de aprendizaje", "Repitencia Escolar", "Asma", "Estreñimiento", "Sudoración en manos"
    ]
    seleccionados = st.multiselect("Checklist de síntomas:", sintomas_list)

# PÁGINA 3-4: HISTORIA FAMILIAR
with tabs[2]:
    st.subheader("VI. INFORMACIÓN FAMILIAR")
    st.write("**Cónyugue**")
    cony_nom = st.text_input("Nombre del Cónyugue:")
    cony_datos = st.text_input("Edad, nivel académico, ocupación y salud:", key="k_cony")
    cony_rel = st.text_area("Describa su relación con su cónyugue e hijos:")
    
    st.write("**Padre**")
    p_nom = st.text_input("Nombre del Padre:")
    p_datos = st.text_input("Edad, Vivo/Muerto, nivel académico, ocupación y salud:", key="k_padre")
    p_rel = st.text_area("Tipo de relación con su padre y castigos recibidos:", key="k_padre_rel")
    
    st.write("**Madre**")
    m_nom = st.text_input("Nombre de la Madre:")
    m_datos = st.text_input("Edad, Vivo/Muerto, nivel académico, ocupación y salud:", key="k_madre")
    m_rel = st.text_area("Tipo de relación con su madre y castigos recibidos:", key="k_madre_rel")
    
    st.divider()
    est_civil_p = st.text_input("Estado civil de los padres (Si hay separación especifique motivo):")
    hermanos = st.text_input("Número de hermanos (varones/mujeres) y lugar que ocupa:")
    h_mejor = st.text_input("¿Con qué hermano se lleva mejor y por qué?")
    crianza = st.text_input("¿Quién fue el encargado de su crianza? (Especifique parentesco):")
    favorito = st.text_input("¿Cree que sus padres tienen un favorito? (¿Quién?):")
    p_opinion = st.text_area("¿Qué opina usted de sus padres? ¿Son cristianos?")
    ant_familiares = st.text_input("Antecedentes familiares (Alcoholismo, maltrato, depresión, enf. mentales):")
    hist_feliz = st.text_area("Cuénteme una historia feliz vivida en familia:")

# PÁGINA 4-5: SOCIAL Y HÁBITOS
with tabs[3]:
    st.subheader("VII-IX. SOCIAL, AMBIENTAL Y HÁBITOS")
    trabajo = st.text_area("Conducta en escuela o trabajo, satisfacción, estrés y cambios laborales:")
    econ = st.select_slider("Situación económica actual:", ["Mala", "Regular", "Buena", "Muy Buena"])
    ley = st.text_input("¿Ha tenido dificultades con la ley? (Especifique):")
    catastrofe = st.text_input("¿Ha sufrido catástrofes naturales o guerras?")
    
    st.write("**Área de Hábitos**")
    c19, c20, c21 = st.columns(3)
    fuma = c19.text_input("Fuma (¿Cuántos al día?):")
    alcohol = c20.text_input("Ingiere bebidas alcohólicas (Especifique):")
    drogas_area = c21.text_input("Consumo de drogas:")
    judicial = st.text_input("Antecedentes judiciales (acusado, detenido, preso):")

# PÁGINA 5: DESARROLLO Y EDUCACIÓN
with tabs[4]:
    st.subheader("X. ANTECEDENTES DEL DESARROLLO")
    embarazo = st.text_input("Circunstancia del embarazo (Planificado / reacción de los padres):")
    parto = st.text_input("Tipo de parto: Recién nacido (uso de fórceps, incubadora):")
    lactancia = st.text_input("Lactancia (natural o artificial, tiempos y motivos):")
    motor = st.text_input("Desarrollo motor (sentarse, gatear, caminar):")
    esfinter = st.text_input("Control de esfínteres, conductas elaboradas de los padres:")
    
    st.write("**Historia Escolar**")
    c22, c23, c24 = st.columns(3)
    esc_ini = c22.text_input("¿Edad cuando entró a la escuela?")
    esc_dur = c23.text_input("¿Tiempo en que la cursó?")
    col_ini = c24.text_input("¿Edad cuando entró al colegio?")
    
    esc_prob = st.text_area("Especifique si tuvo algún problema en su tiempo escolar:")
    c25, c26 = st.columns(2)
    esc_dif = c25.text_input("¿Materias que se le dificultaron más?")
    esc_pref = c26.text_input("¿Materias preferidas?")
    
    tiempo_libre = st.text_input("¿Actividades a las que se dedica en su tiempo libre?")
    repitencia = st.text_input("¿Repitió algún año en su vida escolar?")
    aprender = st.text_input("¿Cómo aprende más fácilmente?")

    st.write("**Historia Sexual**")
    sex_info = st.text_input("Información sexual adquirida / Enfermedades venéreas:")
    menarquia = st.text_input("Periodo menstrual / primera eyaculación (Síntomas):")
    noviazgo_edad = st.text_input("¿Edad del primer noviazgo?")
    noviazgo_opi = st.text_input("Opinión del noviazgo y matrimonio:")
    sex_primera = st.text_input("¿Edad de la primera relación sexual?")
    sex_opi = st.text_area("Opinión de masturbación y relaciones sexuales:")
    homo_opi = st.text_input("Opinión de homosexualidad:")

# PÁGINA 6: PERSONALIDAD Y CIERRE
with tabs[5]:
    st.subheader("XI. PERSONALIDAD PREVIA")
    c27, c28 = st.columns(2)
    pers_1 = c27.text_input("Seguridad en sí mismo:")
    pers_2 = c28.text_input("Toma de decisiones:")
    pers_3 = c27.text_input("Miedo al abandono:")
    pers_4 = c28.text_input("Confianza en otros:")
    pers_5 = c27.text_input("Rencoroso:")
    pers_6 = c28.text_input("Actos impulsivos:")
    pers_7 = st.text_input("Preocupación por el rechazo, crítica, fracaso o ser atractivo:")

    st.subheader("XII. ANÁLISIS Y OBSERVACIONES")
    opinion_prof = st.text_area("OBSERVACIONES GENERALES DEL PSICÓLOGO:")
    
    if st.button("🧠 GENERAR ANÁLISIS TÉCNICO COMPLETO"):
        analisis = motor_ia_dsp({"motivo": motivo, "opinion": opinion_prof, "checks": seleccionados})
        st.session_state["dsp_analisis"] = analisis

    if "dsp_analisis" in st.session_state:
        a = st.session_state["dsp_analisis"]
        st.divider()
        st.info(f"**IMPRESIÓN DIAGNÓSTICA:** {a['estado']}")
        
        col_r, col_t = st.columns(2)
        with col_r:
            st.success(f"**RECOMENDACIÓN:**\n{a['rec']}")
            st.write(f"**Justificación:** {a['just_rec']}")
        with col_t:
            st.warning(f"**BATERÍA DE TESTS:**\n{a['test']}")
            st.write(f"**Justificación:** {a['just_test']}")

    psicologo_firma = st.text_input("Nombre del Psicólogo Evaluador:")
    fecha_app = st.date_input("Fecha de Aplicación:", value=date.today())

# --- 4. EXPORTACIÓN ---
if st.button("💾 GUARDAR Y DESCARGAR EXPEDIENTE"):
    if nombre and identidad:
        doc = Document()
        doc.add_heading('PROTOCOLO DE ENTREVISTA ADULTOS - D.S.P.', 0)
        
        # Resumen del Análisis en el Word
        if "dsp_analisis" in st.session_state:
            a = st.session_state["dsp_analisis"]
            doc.add_heading('ANÁLISIS CLÍNICO IA', level=1)
            doc.add_paragraph(f"Conclusión: {a['estado']}")
            doc.add_paragraph(f"Recomendación Profesional: {a['rec']}")
            doc.add_paragraph(f"Justificación Clínica: {a['just_rec']}")
            doc.add_paragraph(f"Tests Sugeridos: {a['test']}")
            doc.add_paragraph(f"Justificación de Batería: {a['just_test']}")
            
        doc.add_heading('DATOS DEL PACIENTE', level=1)
        doc.add_paragraph(f"Nombre: {nombre}\nIdentidad: {identidad}\nMotivo: {motivo}")
        
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button("📥 DESCARGAR INFORME COMPLETO (WORD)", buf, f"Expediente_{identidad}.docx")
    else:
        st.error("Por favor, ingrese Nombre e Identidad antes de exportar.")
